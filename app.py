#!/usr/bin/env python3
"""
Content Processing Web Application with AI Compliance Analysis

This script combines the robust content extraction and chunking with
improved parallel AI-powered YMYL compliance analysis capabilities.
"""

import streamlit as st
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, WebDriverException
import json
import time
import html
from datetime import datetime
import pytz
import platform
import logging
import asyncio
import aiohttp
from openai import OpenAI
from openai.types.beta.threads import Run
import io
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import markdown
import re
import random
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from enum import Enum

# --- Logging Configuration ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Streamlit Page Configuration ---
st.set_page_config(
    page_title="YMYL Audit Tool",
    page_icon="🚀",
    layout="wide",
)

# --- AI Processing Configuration ---
ANALYZER_ASSISTANT_ID = "asst_WzODK9EapCaZoYkshT6x9xEH"

# --- Enhanced AI Processing Classes ---
class AssistantStatus(Enum):
    SUCCESS = "success"
    RATE_LIMITED = "rate_limited"
    CONTENT_TOO_LONG = "content_too_long"
    ASSISTANT_ERROR = "assistant_error"
    NETWORK_ERROR = "network_error"
    TIMEOUT = "timeout"
    UNKNOWN_ERROR = "unknown_error"

@dataclass
class AssistantResult:
    success: bool
    content: Optional[str] = None
    error: Optional[str] = None
    chunk_index: int = 0
    status: AssistantStatus = AssistantStatus.UNKNOWN_ERROR
    processing_time: float = 0.0
    tokens_used: Optional[int] = None
    retry_count: int = 0

class ImprovedAssistantClient:
    def __init__(self, api_key: str, assistant_id: str, max_concurrent: int = 3):
        self.client = OpenAI(api_key=api_key)
        self.assistant_id = assistant_id
        self.semaphore = asyncio.Semaphore(max_concurrent)
        self.max_content_tokens = 30000  # Token-based limit (safer than character count)
        self.max_retries = 3
        self.base_delay = 1.0
        self.max_delay = 60.0
        
    async def validate_assistant(self) -> bool:
        """Validate that the assistant exists and is accessible."""
        try:
            assistant = self.client.beta.assistants.retrieve(self.assistant_id)
            logger.info(f"Assistant validated: {assistant.name}")
            return True
        except Exception as e:
            logger.error(f"Assistant validation failed: {e}")
            return False

    def _calculate_backoff_delay(self, attempt: int, base_delay: float = None) -> float:
        """Calculate exponential backoff delay with jitter."""
        if base_delay is None:
            base_delay = self.base_delay
            
        # Exponential backoff: base_delay * (2 ^ attempt) + jitter
        delay = min(base_delay * (2 ** attempt), self.max_delay)
        # Add jitter (random factor between 0.5 and 1.5)
        jitter = random.uniform(0.5, 1.5)
        return delay * jitter

    def _validate_content_length(self, content: str) -> bool:
        """Check if content token count is within API limits."""
        # Rough estimate: 4 characters ≈ 1 token (more accurate than character count)
        estimated_tokens = len(content) / 4
        return estimated_tokens <= self.max_content_tokens

    def _map_openai_status_to_assistant_status(self, openai_status: str) -> AssistantStatus:
        """Map OpenAI run status to our AssistantStatus enum."""
        status_mapping = {
            'completed': AssistantStatus.SUCCESS,
            'failed': AssistantStatus.ASSISTANT_ERROR,
            'expired': AssistantStatus.TIMEOUT,
            'cancelled': AssistantStatus.ASSISTANT_ERROR,
            'incomplete': AssistantStatus.CONTENT_TOO_LONG,
            'requires_action': AssistantStatus.ASSISTANT_ERROR,  # Not expected in our use case
        }
        return status_mapping.get(openai_status, AssistantStatus.UNKNOWN_ERROR)
        """Safely extract the latest assistant response from thread."""
        try:
            messages = self.client.beta.threads.messages.list(
                thread_id=thread_id,
                order="desc",
                limit=10  # Get recent messages to find assistant response
            )
            
            # Find the most recent assistant message
            for message in messages.data:
                if message.role == "assistant":
                    if message.content and len(message.content) > 0:
                        return message.content[0].text.value
            
            logger.warning(f"No assistant response found in thread {thread_id}")
            return None
            
        except Exception as e:
            logger.error(f"Error extracting response from thread {thread_id}: {e}")
            return None

    async def _wait_for_run_completion(self, thread_id: str, run_id: str, timeout: int = 300) -> Tuple[bool, Optional[Run]]:
        """Wait for run completion with proper timeout and error handling.
        NOTE: This method is now deprecated in favor of create_and_poll() but kept for compatibility.
        """
        start_time = time.time()
        attempt = 0
        
        while time.time() - start_time < timeout:
            try:
                run = self.client.beta.threads.runs.retrieve(
                    thread_id=thread_id,
                    run_id=run_id
                )
                
                if run.status == 'completed':
                    return True, run
                elif run.status in ['failed', 'cancelled', 'expired']:
                    logger.error(f"Run {run_id} failed with status: {run.status}")
                    return False, run
                elif run.status in ['queued', 'in_progress']:
                    # Use exponential backoff for polling
                    delay = self._calculate_backoff_delay(attempt, 1.0)
                    await asyncio.sleep(min(delay, 10.0))  # Cap at 10 seconds
                    attempt += 1
                else:
                    logger.warning(f"Unknown run status: {run.status}")
                    await asyncio.sleep(2.0)
                    
            except Exception as e:
                logger.error(f"Error checking run status: {e}")
                await asyncio.sleep(5.0)
                
        logger.error(f"Run {run_id} timed out after {timeout} seconds")
        return False, None

    async def _process_single_chunk_with_retries(self, content: str, chunk_index: int) -> AssistantResult:
        """Process a single chunk with retry logic and comprehensive error handling."""
        start_time = time.time()
        
        # Validate content length
        if not self._validate_content_length(content):
            estimated_tokens = len(content) / 4
            return AssistantResult(
                success=False,
                error=f"Content too long: ~{estimated_tokens:.0f} tokens (max: {self.max_content_tokens})",
                chunk_index=chunk_index,
                status=AssistantStatus.CONTENT_TOO_LONG,
                processing_time=time.time() - start_time
            )
        
        last_error = None
        
        for attempt in range(self.max_retries):
            try:
                # Create thread
                thread = self.client.beta.threads.create()
                logger.debug(f"Created thread {thread.id} for chunk {chunk_index}, attempt {attempt + 1}")
                
                try:
                    # Add message to thread
                    self.client.beta.threads.messages.create(
                        thread_id=thread.id,
                        role="user",
                        content=content
                    )
                    
                    # Create and start run
                    run = self.client.beta.threads.runs.create(
                        thread_id=thread.id,
                        assistant_id=self.assistant_id
                    )
                    
                    # Wait for completion
                    completed, final_run = await self._wait_for_run_completion(thread.id, run.id)
                    
                    if completed and final_run:
                        # Extract response
                        response_content = self._extract_assistant_response(thread.id)
                        
                        if response_content:
                            processing_time = time.time() - start_time
                            
                            # Extract token usage if available
                            tokens_used = None
                            if hasattr(final_run, 'usage') and final_run.usage:
                                tokens_used = getattr(final_run.usage, 'total_tokens', None)
                            
                            logger.info(f"Successfully processed chunk {chunk_index} in {processing_time:.2f}s (attempt {attempt + 1})")
                            
                            return AssistantResult(
                                success=True,
                                content=response_content,
                                chunk_index=chunk_index,
                                status=AssistantStatus.SUCCESS,
                                processing_time=processing_time,
                                tokens_used=tokens_used,
                                retry_count=attempt
                            )
                        else:
                            last_error = "No response content found"
                    else:
                        last_error = f"Run failed or timed out: {final_run.status if final_run else 'timeout'}"
                
                finally:
                    # Clean up thread
                    try:
                        self.client.beta.threads.delete(thread.id)
                        logger.debug(f"Cleaned up thread {thread.id}")
                    except Exception as cleanup_error:
                        logger.warning(f"Failed to cleanup thread {thread.id}: {cleanup_error}")
                
            except Exception as e:
                last_error = str(e)
                error_msg = str(e).lower()
                
                # Enhanced error classification with OpenAI-specific errors
                if "rate limit" in error_msg or "429" in error_msg:
                    mapped_status = AssistantStatus.RATE_LIMITED
                    delay = self._calculate_backoff_delay(attempt, 10.0)
                elif "timeout" in error_msg or "connection" in error_msg:
                    mapped_status = AssistantStatus.NETWORK_ERROR
                    delay = self._calculate_backoff_delay(attempt, 2.0)
                elif "maximum context length" in error_msg or "too long" in error_msg:
                    mapped_status = AssistantStatus.CONTENT_TOO_LONG
                    delay = self._calculate_backoff_delay(attempt)
                else:
                    mapped_status = AssistantStatus.ASSISTANT_ERROR
                    delay = self._calculate_backoff_delay(attempt)
                
                logger.warning(f"Chunk {chunk_index} attempt {attempt + 1} failed: {e}")
                
                # Don't wait on the last attempt
                if attempt < self.max_retries - 1:
                    logger.info(f"Retrying chunk {chunk_index} in {delay:.2f} seconds...")
                    await asyncio.sleep(delay)
        
        # All retries exhausted - use the last mapped status
        processing_time = time.time() - start_time
        final_status = mapped_status if 'mapped_status' in locals() else AssistantStatus.ASSISTANT_ERROR
        logger.error(f"Failed to process chunk {chunk_index} after {self.max_retries} attempts: {last_error}")
        
        return AssistantResult(
            success=False,
            error=f"Failed after {self.max_retries} attempts: {last_error}",
            chunk_index=chunk_index,
            status=final_status,
            processing_time=processing_time,
            retry_count=self.max_retries
        )

    async def process_chunk(self, content: str, chunk_index: int) -> AssistantResult:
        """Process a single chunk with semaphore-controlled concurrency."""
        async with self.semaphore:
            return await self._process_single_chunk_with_retries(content, chunk_index)

    async def process_chunks_parallel(self, chunks: List[Dict]) -> List[AssistantResult]:
        """Process multiple chunks in parallel with proper resource management."""
        try:
            # Validate assistant first
            if not await self.validate_assistant():
                return [AssistantResult(
                    success=False,
                    error="Assistant validation failed",
                    chunk_index=0,
                    status=AssistantStatus.ASSISTANT_ERROR
                )]
            
            logger.info(f"Starting parallel processing of {len(chunks)} chunks with max concurrency: {self.semaphore._value}")
            
            # Create tasks for all chunks
            tasks = []
            for chunk in chunks:
                task = asyncio.create_task(
                    self.process_chunk(chunk["text"], chunk["index"])
                )
                tasks.append(task)
            
            # Wait for all tasks to complete
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process results and handle exceptions
            processed_results = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    processed_results.append(AssistantResult(
                        success=False,
                        error=f"Task exception: {str(result)}",
                        chunk_index=chunks[i]["index"],
                        status=AssistantStatus.UNKNOWN_ERROR
                    ))
                else:
                    processed_results.append(result)
            
            # Log summary statistics
            successful = sum(1 for r in processed_results if r.success)
            failed = len(processed_results) - successful
            total_time = max((r.processing_time for r in processed_results), default=0)
            
            logger.info(f"Parallel processing completed: {successful}/{len(chunks)} successful in {total_time:.2f}s")
            
            return processed_results
            
        except Exception as e:
            logger.error(f"Error in parallel processing: {e}")
            return [AssistantResult(
                success=False,
                error=f"Parallel processing failed: {str(e)}",
                chunk_index=0,
                status=AssistantStatus.UNKNOWN_ERROR
            )]

# --- Component 1: Updated Content Extractor ---
class ContentExtractor:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def extract_content(self, url):
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            
            content_parts = []
            
            # 1. Extract H1 (anywhere on page)
            h1 = soup.find('h1')
            if h1:
                text = h1.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"H1: {text}")
            
            # 2. Extract Subtitle (anywhere on page)
            subtitle = soup.find('span', class_=['sub-title', 'd-block'])
            if subtitle:
                text = subtitle.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"SUBTITLE: {text}")
            
            # 3. Extract Lead (anywhere on page)
            lead = soup.find('p', class_='lead')
            if lead:
                text = lead.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"LEAD: {text}")
            
            # 4. Extract Article content
            article = soup.find('article')
            if article:
                # Remove tab-content sections before processing
                for tab_content in article.find_all('div', class_='tab-content'):
                    tab_content.decompose()
                
                # Process all elements in document order within article
                for element in article.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'span', 'p']):
                    text = element.get_text(separator='\n', strip=True)
                    if not text:
                        continue
                    
                    # Check element type and add appropriate prefix
                    if element.name == 'h1':
                        content_parts.append(f"H1: {text}")
                    elif element.name == 'h2':
                        content_parts.append(f"H2: {text}")
                    elif element.name == 'h3':
                        content_parts.append(f"H3: {text}")
                    elif element.name == 'h4':
                        content_parts.append(f"H4: {text}")
                    elif element.name == 'h5':
                        content_parts.append(f"H5: {text}")
                    elif element.name == 'h6':
                        content_parts.append(f"H6: {text}")
                    elif element.name == 'span' and 'sub-title' in element.get('class', []) and 'd-block' in element.get('class', []):
                        content_parts.append(f"SUBTITLE: {text}")
                    elif element.name == 'p' and 'lead' in element.get('class', []):
                        content_parts.append(f"LEAD: {text}")
                    elif element.name == 'p':
                        content_parts.append(f"CONTENT: {text}")
            
            # 5. Extract FAQ section
            faq_section = soup.find('section', attrs={'data-qa': 'templateFAQ'})
            if faq_section:
                text = faq_section.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"FAQ: {text}")
            
            # 6. Extract Author section
            author_section = soup.find('section', attrs={'data-qa': 'templateAuthorCard'})
            if author_section:
                text = author_section.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"AUTHOR: {text}")
            
            # Join with double newlines to preserve spacing
            final_content = '\n\n'.join(content_parts)
            return True, final_content, None
            
        except requests.RequestException as e:
            return False, None, f"Error fetching URL: {e}"
        except Exception as e:
            return False, None, f"Error processing content: {e}"

# --- Component 2: The Final, Upgraded Chunk Processor ---
class ChunkProcessor:
    def __init__(self, log_callback=None):
        self.driver = None
        self.log = log_callback if log_callback else logger.info

    def _setup_driver(self):
        self.log("Initializing browser with enhanced stability & permissions...")
        chrome_options = Options()
        chrome_options.add_argument('--headless=new')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_experimental_option("prefs", {"profile.default_content_setting_values.clipboard": 1})
        try:
            self.driver = webdriver.Chrome(options=chrome_options)
            self.log("✅ Browser initialized successfully.")
            return True
        except WebDriverException as e:
            self.log(f"❌ WebDriver Initialization Failed: {e}")
            return False

    def _extract_json_from_button(self):
        try:
            wait = WebDriverWait(self.driver, 180)
            h3_xpath = "//h3[text()='Raw JSON Output']"
            self.log("🔄 Waiting for results section to appear...")
            wait.until(EC.presence_of_element_located((By.XPATH, h3_xpath)))
            self.log("✅ Results section is visible.")
            button_selector = "button[data-testid='stCodeCopyButton']"
            self.log("...Waiting for the copy button...")
            copy_button = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, button_selector)))
            self.log("✅ Found the copy button element.")
            self.log("...Polling button's attribute for completeness...")
            timeout = time.time() + 10
            final_content = ""
            while time.time() < timeout:
                raw_content = copy_button.get_attribute('data-clipboard-text')
                if raw_content and raw_content.strip().startswith('{') and raw_content.strip().endswith('}'):
                    final_content = raw_content; break
                time.sleep(0.2)
            if not final_content: self.log("❌ Timed out polling the attribute."); return None
            self.log("...Decoding HTML entities...")
            decoded_content = html.unescape(final_content)
            self.log(f"✅ Extraction complete. Retrieved {len(decoded_content):,} characters.")
            return decoded_content
        except Exception as e:
            self.log(f"❌ An error occurred during the final JSON extraction phase: {e}")
            return None

    def process_content(self, content):
        if not self._setup_driver():
            return False, None, "Failed to initialize browser."
        try:
            self.log(f"Navigating to `chunk.dejan.ai`...")
            self.driver.get("https://chunk.dejan.ai/")
            wait = WebDriverWait(self.driver, 30)
            self.log("Using JavaScript to copy full text to browser's clipboard...")
            self.driver.execute_script("navigator.clipboard.writeText(arguments[0]);", content)
            self.log("Locating text area and clearing it...")
            textarea_selector = (By.CSS_SELECTOR, 'textarea[aria-label="Text to chunk:"]')
            input_field = wait.until(EC.element_to_be_clickable(textarea_selector))
            input_field.clear()
            self.log("Simulating a 'Paste' (Ctrl+V) command...")
            modifier_key = Keys.COMMAND if platform.system() == "Darwin" else Keys.CONTROL
            input_field.send_keys(modifier_key, "v")
            self.log("Clicking submit button...")
            submit_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '[data-testid="stBaseButton-secondary"]')))
            submit_button.click()
            json_output = self._extract_json_from_button()
            if json_output:
                return True, json_output, None
            else:
                return False, None, "Failed to extract JSON from the results page."
        except Exception as e:
            return False, None, f"An unexpected error occurred during processing: {e}"
        finally:
            self.cleanup()
            
    def cleanup(self):
        if self.driver:
            self.log("Cleaning up and closing browser instance.")
            self.driver.quit()
            self.log("✅ Browser closed.")

# --- Component 3: Enhanced AI Processing Functions ---

def convert_to_html(markdown_content):
    """Convert markdown report to styled HTML"""
    try:
        # Convert markdown to HTML
        html_content = markdown.markdown(markdown_content, extensions=['tables', 'toc'])
        
        # Add professional CSS styling
        css_style = """
        <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }
        h1 {
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }
        h2 {
            color: #34495e;
            border-left: 4px solid #3498db;
            padding-left: 10px;
            margin-top: 30px;
        }
        h3 {
            color: #34495e;
            margin-top: 25px;
        }
        .severity-critical { color: #e74c3c; font-weight: bold; }
        .severity-high { color: #e67e22; font-weight: bold; }
        .severity-medium { color: #f39c12; font-weight: bold; }
        .severity-low { color: #3498db; font-weight: bold; }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th, td {
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        .processing-summary {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }
        code {
            background-color: #f1f2f6;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }
        blockquote {
            border-left: 4px solid #bdc3c7;
            margin: 0;
            padding-left: 15px;
            color: #7f8c8d;
        }
        </style>
        """
        
        # Enhance severity indicators
        html_content = html_content.replace('🔴', '<span class="severity-critical">🔴</span>')
        html_content = html_content.replace('🟠', '<span class="severity-high">🟠</span>')
        html_content = html_content.replace('🟡', '<span class="severity-medium">🟡</span>')
        html_content = html_content.replace('🔵', '<span class="severity-low">🔵</span>')
        
        # Wrap processing summary
        html_content = re.sub(
            r'## Processing Summary(.*?)(?=##|$)', 
            r'<div class="processing-summary"><h2>Processing Summary</h2>\1</div>', 
            html_content, 
            flags=re.DOTALL
        )
        
        # Complete HTML document
        full_html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>YMYL Compliance Audit Report</title>
            {css_style}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        return full_html.encode('utf-8')
    except Exception as e:
        logger.error(f"HTML conversion error: {e}")
        return f"<html><body><h1>Export Error</h1><p>Failed to convert report: {e}</p></body></html>".encode('utf-8')

def convert_to_word(markdown_content):
    """Convert markdown report to Word document"""
    try:
        doc = Document()
        
        # Set document styles
        styles = doc.styles
        
        # Create custom heading styles
        if 'Report Title' not in styles:
            title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Inches(0.2)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(44, 62, 80)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parse markdown content
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Handle headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.style = 'Report Title'
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('- '):
                # Bullet points
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('---'):
                # Horizontal rule (skip)
                continue
            elif '🔴' in line or '🟠' in line or '🟡' in line or '🔵' in line:
                # Severity indicators - make them stand out
                p = doc.add_paragraph(line)
                if '🔴' in line:
                    p.runs[0].font.color.rgb = RGBColor(231, 76, 60)  # Red
                elif '🟠' in line:
                    p.runs[0].font.color.rgb = RGBColor(230, 126, 34)  # Orange
                elif '🟡' in line:
                    p.runs[0].font.color.rgb = RGBColor(243, 156, 18)  # Yellow/Gold
                elif '🔵' in line:
                    p.runs[0].font.color.rgb = RGBColor(52, 152, 219)  # Blue
                p.runs[0].font.bold = True
            else:
                # Regular paragraph
                if line:
                    doc.add_paragraph(line)
        
        # Save to memory
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Word conversion error: {e}")
        # Return simple document with error message
        doc = Document()
        doc.add_heading('Export Error', 0)
        doc.add_paragraph(f'Failed to convert report: {e}')
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()

def convert_to_pdf(markdown_content):
    """Convert markdown report to PDF document"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#34495e'),
            spaceBefore=12,
            spaceAfter=6
        )
        
        critical_style = ParagraphStyle(
            'Critical',
            parent=styles['Normal'],
            textColor=colors.red,
            fontSize=10,
            fontName='Helvetica-Bold'
        )
        
        high_style = ParagraphStyle(
            'High',
            parent=styles['Normal'],
            textColor=colors.orange,
            fontSize=10,
            fontName='Helvetica-Bold'
        )
        
        # Build story
        story = []
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                story.append(Spacer(1, 12))
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('**') and line.endswith('**'):
                story.append(Paragraph(f"<b>{line[2:-2]}</b>", styles['Normal']))
            elif line.startswith('---'):
                story.append(Spacer(1, 12))
            elif '🔴' in line:
                story.append(Paragraph(line, critical_style))
            elif '🟠' in line:
                story.append(Paragraph(line, high_style))
            elif line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            else:
                if line:
                    story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"PDF conversion error: {e}")
        # Return simple error PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = [
            Paragraph("Export Error", getSampleStyleSheet()['Title']),
            Spacer(1, 12),
            Paragraph(f"Failed to convert report: {e}", getSampleStyleSheet()['Normal'])
        ]
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

def create_export_options(report_content):
    """Create multiple export format options"""
    return {
        'html': convert_to_html(report_content),
        'docx': convert_to_word(report_content),
        'pdf': convert_to_pdf(report_content),
        'markdown': report_content.encode('utf-8')
    }

def extract_big_chunks(json_data):
    """Extract and format big chunks for AI processing."""
    try:
        big_chunks = json_data.get('big_chunks', [])
        chunks = []
        
        for chunk in big_chunks:
            chunk_index = chunk.get('big_chunk_index', len(chunks) + 1)
            small_chunks = chunk.get('small_chunks', [])
            
            # Join small chunks with newlines
            joined_text = '\n'.join(small_chunks)
            
            chunks.append({
                "index": chunk_index,
                "text": joined_text,
                "count": len(small_chunks)
            })
        
        return chunks
    except Exception as e:
        logger.error(f"Error extracting big chunks: {e}")
        return []

def create_final_report_simple(analysis_results):
    """Create final report by simple concatenation."""
    try:
        report_parts = []
        
        # Add header
        audit_date = datetime.now().strftime("%Y-%m-%d")
        header = f"""# YMYL Compliance Audit Report

**Audit Date:** {audit_date}
**Content Type:** Online Casino/Gambling  
**Analysis Method:** Section-by-section E-E-A-T compliance review

---

"""
        report_parts.append(header)
        
        # Add successful analyses
        successful_count = 0
        error_count = 0
        
        for result in analysis_results:
            if result.get("success"):
                report_parts.append(result["content"])
                report_parts.append("\n---\n")
                successful_count += 1
            else:
                error_count += 1
                error_section = f"""
# Analysis Error for Chunk {result.get('chunk_index', 'Unknown')}

❌ **Processing Failed**
Error: {result.get('error', 'Unknown error')}

---
"""
                report_parts.append(error_section)
        
        # Add processing summary
        total_sections = successful_count + error_count
        summary = f"""
## Processing Summary
**✅ Sections Successfully Analyzed:** {successful_count}
**❌ Sections with Analysis Errors:** {error_count}  
**📊 Total Sections:** {total_sections}

---
*Report generated by AI-powered YMYL compliance analysis system*
"""
        report_parts.append(summary)
        
        return ''.join(report_parts)
        
    except Exception as e:
        logger.error(f"Error creating final report: {e}")
        return f"Error generating report: {e}"

# --- Enhanced AI Analysis Function ---
async def process_ai_analysis_improved(json_output, api_key, log_callback=None):
    """Improved AI compliance analysis with better error handling and logging."""
    def log(message):
        if log_callback: 
            log_callback(message)
        logger.info(message)
    
    try:
        # Parse JSON and extract chunks
        log("📊 Parsing JSON and extracting chunks...")
        json_data = json.loads(json_output)
        chunks = extract_big_chunks(json_data)
        
        if not chunks:
            return False, "No chunks found in JSON data", None
            
        log(f"🚀 Starting improved parallel analysis of {len(chunks)} chunks...")
        
        # Create improved assistant client
        assistant_client = ImprovedAssistantClient(
            api_key=api_key,
            assistant_id=ANALYZER_ASSISTANT_ID,
            max_concurrent=3  # Adjust based on your rate limits
        )
        
        # Process chunks
        results = await assistant_client.process_chunks_parallel(chunks)
        
        # Convert to format expected by existing code
        analysis_results = []
        for result in results:
            analysis_results.append({
                "success": result.success,
                "content": result.content,
                "error": result.error,
                "chunk_index": result.chunk_index,
                "tokens_used": result.tokens_used,
                "processing_time": result.processing_time,
                "retry_count": result.retry_count,
                "status": result.status.value
            })
        
        # Create final report using existing function
        log("📝 Assembling final report...")
        final_report = create_final_report_simple(analysis_results)
        
        log("🎉 Improved AI Analysis Complete!")
        return True, final_report, analysis_results
        
    except json.JSONDecodeError as e:
        error_msg = f"Invalid JSON format: {e}"
        log(f"❌ {error_msg}")
        return False, error_msg, None
    except Exception as e:
        error_msg = f"AI analysis error: {e}"
        log(f"❌ {error_msg}")
        return False, error_msg, None

async def process_ai_analysis_improved_with_settings(json_output, api_key, max_concurrent=3, max_retries=3, max_tokens=30000, log_callback=None):
    """Improved AI compliance analysis with custom settings."""
    def log(message):
        if log_callback: 
            log_callback(message)
        logger.info(message)
    
    try:
        # Parse JSON and extract chunks
        log("📊 Parsing JSON and extracting chunks...")
        json_data = json.loads(json_output)
        chunks = extract_big_chunks(json_data)
        
        if not chunks:
            return False, "No chunks found in JSON data", None
            
        log(f"🚀 Starting improved parallel analysis of {len(chunks)} chunks...")
        
        # Create improved assistant client with custom settings
        assistant_client = ImprovedAssistantClient(
            api_key=api_key,
            assistant_id=ANALYZER_ASSISTANT_ID,
            max_concurrent=max_concurrent
        )
        assistant_client.max_retries = max_retries
        assistant_client.max_content_tokens = max_tokens
        
        # Process chunks
        results = await assistant_client.process_chunks_parallel(chunks)
        
        # Convert to format expected by existing code
        analysis_results = []
        for result in results:
            analysis_results.append({
                "success": result.success,
                "content": result.content,
                "error": result.error,
                "chunk_index": result.chunk_index,
                "tokens_used": result.tokens_used,
                "processing_time": result.processing_time,
                "retry_count": result.retry_count,
                "status": result.status.value
            })
        
        # Create final report using existing function
        log("📝 Assembling final report...")
        final_report = create_final_report_simple(analysis_results)
        
        log("🎉 Improved AI Analysis Complete!")
        return True, final_report, analysis_results
        
    except json.JSONDecodeError as e:
        error_msg = f"Invalid JSON format: {e}"
        log(f"❌ {error_msg}")
        return False, error_msg, None
    except Exception as e:
        error_msg = f"AI analysis error: {e}"
        log(f"❌ {error_msg}")
        return False, error_msg, None

# --- Main Workflow Function ---
def process_url_workflow_with_logging(url, log_callback=None):
    result = {'success': False, 'url': url, 'extracted_content': None, 'json_output': None, 'error': None}
    
    def log(message):
        if log_callback: log_callback(message)
        logger.info(message)
        
    try:
        log("🚀 Initializing content extractor...")
        extractor = ContentExtractor()
        log(f"🔍 Fetching and extracting content from: {url}")
        success, content, error = extractor.extract_content(url)
        if not success:
            result['error'] = f"Content extraction failed: {error}"; return result
        result['extracted_content'] = content
        log(f"✅ Content extracted: {len(content):,} characters")

        log("🤖 Initializing chunk processor...")
        processor = ChunkProcessor(log_callback=log)
        success, json_output, error = processor.process_content(content)
        if not success:
            result['error'] = f"Chunk processing failed: {error}"; return result
        
        result['json_output'] = json_output
        result['success'] = True
        log("🎉 Workflow Complete!")
        return result
    except Exception as e:
        log(f"💥 An unexpected error occurred in the workflow: {str(e)}")
        result['error'] = f"An unexpected workflow error occurred: {str(e)}"
        return result

# --- Streamlit UI ---
def main():
    st.title("🕵 YMYL Audit Tool")
    st.markdown("**Automatically extract content from websites, generate JSON chunks, and perform YMYL compliance analysis**")

    # Sidebar configuration
    debug_mode = st.sidebar.checkbox("🐛 Debug Mode", value=True, help="Show detailed processing logs")

    # API Key configuration
    st.sidebar.markdown("### 🔑 AI Analysis Configuration")
    try:
        api_key = st.secrets["openai_api_key"]
        st.sidebar.success("✅ API Key loaded from secrets")
    except Exception:
        api_key = st.sidebar.text_input(
            "OpenAI API Key:",
            type="password",
            help="Enter your OpenAI API key for AI analysis"
        )
        if api_key:
            st.sidebar.success("✅ API Key provided")
        else:
            st.sidebar.warning("⚠️ API Key needed for AI analysis")

    # Enhanced Configuration Options
    st.sidebar.markdown("### ⚙️ Advanced Settings")
    max_concurrent = st.sidebar.slider(
        "Max Concurrent Analyses", 
        min_value=1, 
        max_value=10, 
        value=3,
        help="Number of chunks to process simultaneously (higher = faster but may hit rate limits)"
    )
    
    max_retries = st.sidebar.slider(
        "Max Retries per Chunk",
        min_value=1,
        max_value=5,
        value=3,
        help="Number of retry attempts for failed chunks"
    )
    
    max_tokens = st.sidebar.number_input(
        "Max Content Tokens",
        min_value=5000,
        max_value=50000,
        value=30000,
        step=5000,
        help="Maximum tokens per chunk (4 chars ≈ 1 token)"
    )

    st.markdown("---")
    col1, col2 = st.columns([2, 1])

    with col1:
        url = st.text_input("Enter the URL to process:", help="Include http:// or https://")
        if st.button("🚀 Process URL", type="primary", use_container_width=True):
            if not url:
                st.error("Please enter a URL to process")
                return

            # clear old state
            for key in ("latest_result", "ai_analysis_result"):
                st.session_state.pop(key, None)

            if debug_mode:
                # Detailed logging
                log_placeholder = st.empty()
                log_lines = []
                def log_callback(msg):
                    now = datetime.now(pytz.timezone("Europe/Malta"))
                    log_lines.append(f"`{now.strftime('%H:%M:%S')}`: {msg}")
                    log_placeholder.info("\n".join(log_lines))

                result = process_url_workflow_with_logging(url, log_callback)
                st.session_state["latest_result"] = result
                if result["success"]:
                    st.success("Processing completed successfully!")
                else:
                    st.error(f"Error: {result['error']}")
            else:
                # Simple milestones
                log_area = st.empty()
                milestones = []
                def simple_log(text):
                    milestones.append(f"- {text}")
                    log_area.markdown("\n".join(milestones))

                simple_log("Extracting content")
                extractor = ContentExtractor()
                ok, content, err = extractor.extract_content(url)
                if not ok:
                    st.error(f"Error: {err}")
                    return

                simple_log("Sending content to Chunk Norris")
                processor = ChunkProcessor()

                with st.status("You are not waiting, Chunk Norris is waiting for you"):
                    ok, json_out, err = processor.process_content(content)

                simple_log("Chunking done!")
                st.success("Chunking done!")

                st.session_state["latest_result"] = {
                    "success": ok,
                    "url": url,
                    "extracted_content": content if ok else None,
                    "json_output": json_out if ok else None,
                    "error": err
                }

    with col2:
        st.subheader("ℹ️ How it works")
        st.markdown("""
1. **Extract**: Extract the content.
2. **Chunk**: Send extracted text to Chunk Norris.
3. **YMYL Analysis**: YMYL audit of the content with AI.
4. **Done**: Output complete report.
""")
        st.info("💡 **New**: Enhanced AI analysis with improved error handling and retry logic!")

    # Results Display
    if 'latest_result' in st.session_state and st.session_state['latest_result'].get('success'):
        result = st.session_state['latest_result']
        st.markdown("---")
        st.subheader("📊 Results")
        
        # Enhanced AI Analysis Button
        if api_key and st.button("🤖 Process with Enhanced AI Compliance Analysis", type="secondary", use_container_width=True):
            try:
                # Parse JSON and extract chunks first
                json_data = json.loads(result['json_output'])
                chunks = extract_big_chunks(json_data)
                
                if not chunks:
                    st.error("No chunks found in JSON data")
                    return
                
                # Enhanced Processing Logs Section
                st.subheader("🔍 Enhanced Processing Logs")
                log_container = st.container()
                
                with log_container:
                    st.info(f"🚀 Starting improved parallel analysis of {len(chunks)} chunks...")
                    st.write("**Configuration:**")
                    st.write(f"- Analyzer: `{ANALYZER_ASSISTANT_ID}`")
                    st.write("- Report Maker: `Simple Concatenation (No AI)`")
                    st.write(f"- API Key Status: {'✅ Valid' if api_key.startswith('sk-') else '❌ Invalid'}")
                    st.write(f"- Max Concurrent: `{max_concurrent}`")
                    st.write(f"- Max Retries: `{max_retries}`")
                    st.write(f"- Max Tokens: `{max_tokens:,}`")
                    st.write("**Chunk Details:**")
                    for chunk in chunks:
                        estimated_tokens = len(chunk['text']) / 4
                        st.write(f"- Chunk {chunk['index']}: {len(chunk['text']):,} chars (~{estimated_tokens:.0f} tokens)")
                
                # Progress tracking with enhanced metrics
                total_chunks = len(chunks)
                progress_bar = st.progress(0)
                status_container = st.empty()
                metrics_container = st.container()
                
                # Real-time metrics display
                with metrics_container:
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        processing_metric = st.empty()
                    with col2:
                        success_metric = st.empty()
                    with col3:
                        failure_metric = st.empty()
                    with col4:
                        speed_metric = st.empty()
                
                # Start processing with timing
                start_time = time.time()
                
                with st.spinner("🤖 Running enhanced parallel analysis..."):
                    # Create improved assistant client with user settings
                    assistant_client = ImprovedAssistantClient(
                        api_key=api_key,
                        assistant_id=ANALYZER_ASSISTANT_ID,
                        max_concurrent=max_concurrent
                    )
                    assistant_client.max_retries = max_retries
                    
                    # Run improved AI analysis with custom settings
                    success, ai_result, analysis_details = asyncio.run(process_ai_analysis_improved_with_settings(
                        result['json_output'], 
                        api_key,
                        max_concurrent,
                        max_retries,
                        max_tokens,
                        None  # Disable callback since we have enhanced UI
                    ))
                
                # Update progress
                progress_bar.progress(1.0)
                processing_time = time.time() - start_time
                
                # Display enhanced processing summary
                if success and analysis_details:
                    successful_analyses = [r for r in analysis_details if r.get("success")]
                    failed_analyses = [r for r in analysis_details if not r.get("success")]
                    
                    # Calculate additional metrics
                    total_tokens = sum(r.get('tokens_used', 0) for r in successful_analyses if r.get('tokens_used'))
                    avg_processing_time = sum(r.get('processing_time', 0) for r in analysis_details) / len(analysis_details)
                    total_retries = sum(r.get('retry_count', 0) for r in analysis_details)
                    
                    with status_container.container():
                        # Enhanced metrics display
                        col1, col2, col3, col4, col5 = st.columns(5)
                        with col1:
                            st.metric("Total Chunks", total_chunks)
                        with col2:
                            st.metric("Successful", len(successful_analyses), 
                                     delta=len(successful_analyses) if len(successful_analyses) == total_chunks else None)
                        with col3:
                            st.metric("Failed", len(failed_analyses), 
                                     delta=f"-{len(failed_analyses)}" if len(failed_analyses) > 0 else None)
                        with col4:
                            st.metric("Total Tokens", f"{total_tokens:,}" if total_tokens > 0 else "N/A")
                        with col5:
                            st.metric("Total Retries", total_retries)
                        
                        # Performance metrics
                        col6, col7, col8 = st.columns(3)
                        with col6:
                            st.metric("Processing Time", f"{processing_time:.2f}s")
                        with col7:
                            st.metric("Avg Time/Chunk", f"{avg_processing_time:.2f}s")
                        with col8:
                            success_rate = (len(successful_analyses) / total_chunks * 100) if total_chunks > 0 else 0
                            st.metric("Success Rate", f"{success_rate:.1f}%")
                        
                        # Status indicator
                        if len(failed_analyses) == 0:
                            st.success(f"✅ All chunks processed successfully in {processing_time:.2f} seconds")
                        elif len(successful_analyses) > 0:
                            st.warning(f"⚠️ Partial success: {len(successful_analyses)}/{total_chunks} chunks processed in {processing_time:.2f} seconds")
                        else:
                            st.error(f"❌ All chunks failed to process in {processing_time:.2f} seconds")
                    
                    # Store enhanced results
                    st.session_state['ai_analysis_result'] = {
                        'success': True,
                        'report': ai_result,
                        'details': analysis_details,
                        'processing_time': processing_time,
                        'total_chunks': total_chunks,
                        'successful_count': len(successful_analyses),
                        'failed_count': len(failed_analyses),
                        'total_tokens': total_tokens,
                        'avg_processing_time': avg_processing_time,
                        'total_retries': total_retries,
                        'success_rate': success_rate
                    }
                    
                else:
                    st.session_state['ai_analysis_result'] = {
                        'success': False,
                        'error': ai_result if not success else 'Unknown error occurred'
                    }
                    st.error(f"❌ Enhanced AI analysis failed: {ai_result if not success else 'Unknown error'}")
                    
            except json.JSONDecodeError as e:
                st.error(f"❌ Invalid JSON format: {str(e)}")
            except Exception as e:
                st.error(f"❌ An error occurred during enhanced AI analysis: {str(e)}")
                logger.error(f"Enhanced AI analysis error: {str(e)}")

        # Enhanced Results Tabs
        if 'ai_analysis_result' in st.session_state and st.session_state['ai_analysis_result'].get('success'):
            tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
                "🎯 AI Compliance Report", 
                "📊 Individual Analyses", 
                "📈 Performance Metrics",
                "🔧 JSON Output", 
                "📄 Extracted Content", 
                "📋 Summary"
            ])
            
            with tab1:
                st.markdown("### YMYL Compliance Analysis Report")
                ai_report = st.session_state['ai_analysis_result']['report']
                
                # Enhanced Export Options
                st.markdown("#### 📋 Copy Report")
                st.code(ai_report, language='markdown')
                
                # Multiple Export Format Options
                st.markdown("#### 📄 Download Formats")
                st.markdown("Choose your preferred format for professional use:")
                
                # Create export data
                try:
                    export_formats = create_export_options(ai_report)
                    timestamp = int(time.time())
                    
                    # Create download buttons in columns
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.download_button(
                            label="📝 Markdown",
                            data=export_formats['markdown'],
                            file_name=f"ymyl_compliance_report_{timestamp}.md",
                            mime="text/markdown",
                            help="Original markdown format - perfect for copying to other platforms"
                        )
                    
                    with col2:
                        st.download_button(
                            label="🌐 HTML",
                            data=export_formats['html'],
                            file_name=f"ymyl_compliance_report_{timestamp}.html",
                            mime="text/html",
                            help="Styled HTML document - opens in any web browser"
                        )
                    
                    with col3:
                        st.download_button(
                            label="📄 Word",
                            data=export_formats['docx'],
                            file_name=f"ymyl_compliance_report_{timestamp}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            help="Microsoft Word document - ready for editing and sharing"
                        )
                    
                    with col4:
                        st.download_button(
                            label="📋 PDF",
                            data=export_formats['pdf'],
                            file_name=f"ymyl_compliance_report_{timestamp}.pdf",
                            mime="application/pdf",
                            help="Professional PDF document - perfect for presentations and archival"
                        )
                    
                    st.info("""
                    💡 **Format Guide:**
                    - **Markdown**: Best for developers and copy-pasting to other platforms
                    - **HTML**: Opens in web browsers, styled and formatted
                    - **Word**: Professional business format, editable and shareable
                    - **PDF**: Final presentation format, preserves formatting across devices
                    """)
                    
                except Exception as e:
                    st.error(f"Error creating export formats: {e}")
                    # Fallback to basic markdown download
                    st.download_button(
                        label="💾 Download Report (Markdown)",
                        data=ai_report,
                        file_name=f"ymyl_compliance_report_{timestamp}.md",
                        mime="text/markdown"
                    )
                
                with st.expander("📖 View Formatted Report"):
                    st.markdown(ai_report)
            
            with tab2:
                st.markdown("### Individual Chunk Analysis Results")
                analysis_details = st.session_state['ai_analysis_result']['details']
                
                # Enhanced processing metrics at top
                ai_result = st.session_state['ai_analysis_result']
                col1, col2, col3, col4, col5 = st.columns(5)
                with col1:
                    st.metric("Processing Time", f"{ai_result.get('processing_time', 0):.2f}s")
                with col2:
                    st.metric("Total Chunks", ai_result.get('total_chunks', 0))
                with col3:
                    st.metric("Successful", ai_result.get('successful_count', 0))
                with col4:
                    st.metric("Failed", ai_result.get('failed_count', 0))
                with col5:
                    st.metric("Success Rate", f"{ai_result.get('success_rate', 0):.1f}%")
                
                st.markdown("---")
                
                # Enhanced individual chunk results with status indicators
                for detail in analysis_details:
                    chunk_idx = detail.get('chunk_index', 'Unknown')
                    processing_time = detail.get('processing_time', 0)
                    retry_count = detail.get('retry_count', 0)
                    status = detail.get('status', 'unknown')
                    tokens_used = detail.get('tokens_used')
                    
                    if detail.get('success'):
                        status_emoji = "✅"
                        status_text = "Success"
                        expand_default = False
                    else:
                        status_emoji = "❌" 
                        status_text = f"Failed ({status})"
                        expand_default = True
                    
                    # Create enhanced header with metrics
                    header = f"{status_emoji} Chunk {chunk_idx} - {status_text}"
                    if processing_time > 0:
                        header += f" ({processing_time:.2f}s"
                        if retry_count > 0:
                            header += f", {retry_count} retries"
                        if tokens_used:
                            header += f", {tokens_used} tokens"
                        header += ")"
                    
                    with st.expander(header, expanded=expand_default):
                        if detail.get('success'):
                            st.markdown(detail['content'])
                        else:
                            st.error(f"**Error**: {detail.get('error', 'Unknown error')}")
                            st.write(f"**Status**: {status}")
                            if retry_count > 0:
                                st.write(f"**Retry attempts**: {retry_count}")
            
            with tab3:
                st.markdown("### Performance Metrics & Analytics")
                ai_result = st.session_state['ai_analysis_result']
                analysis_details = ai_result['details']
                
                # Performance overview
                st.markdown("#### Overall Performance")
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Total Processing Time", f"{ai_result.get('processing_time', 0):.2f}s")
                with col2:
                    st.metric("Average Time per Chunk", f"{ai_result.get('avg_processing_time', 0):.2f}s")
                with col3:
                    st.metric("Total API Tokens Used", f"{ai_result.get('total_tokens', 0):,}")
                with col4:
                    st.metric("Total Retry Attempts", ai_result.get('total_retries', 0))
                
                # Status breakdown
                st.markdown("#### Status Breakdown")
                status_counts = {}
                for detail in analysis_details:
                    status = detail.get('status', 'unknown')
                    status_counts[status] = status_counts.get(status, 0) + 1
                
                status_cols = st.columns(len(status_counts))
                for i, (status, count) in enumerate(status_counts.items()):
                    with status_cols[i]:
                        st.metric(f"{status.replace('_', ' ').title()}", count)
                
                # Performance insights
                st.markdown("#### Performance Insights")
                if ai_result.get('processing_time', 0) > 0 and len(analysis_details) > 0:
                    parallel_efficiency = (ai_result.get('avg_processing_time', 0) * len(analysis_details)) / ai_result.get('processing_time', 1)
                    st.info(f"📊 **Parallel Efficiency**: {parallel_efficiency:.1f}x faster than sequential processing")
                
                if ai_result.get('total_retries', 0) > 0:
                    st.warning(f"⚠️ **Reliability**: {ai_result.get('total_retries', 0)} retry attempts were needed")
                else:
                    st.success("✅ **Reliability**: All chunks processed successfully on first attempt")
                
                # Detailed timing analysis
                if len(analysis_details) > 1:
                    st.markdown("#### Timing Analysis")
                    processing_times = [d.get('processing_time', 0) for d in analysis_details]
                    min_time = min(processing_times)
                    max_time = max(processing_times)
                    avg_time = sum(processing_times) / len(processing_times)
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Fastest Chunk", f"{min_time:.2f}s")
                    with col2:
                        st.metric("Slowest Chunk", f"{max_time:.2f}s")
                    with col3:
                        st.metric("Time Variance", f"±{(max_time - min_time):.2f}s")
            
            with tab4:
                st.code(result['json_output'], language='json')
                st.download_button(
                    label="💾 Download JSON",
                    data=result['json_output'],
                    file_name=f"chunks_{int(time.time())}.json",
                    mime="application/json"
                )
            
            with tab5:
                st.text_area("Raw extracted content:", value=result['extracted_content'], height=400)
            
            with tab6:
                st.subheader("Complete Processing Summary")
                try:
                    json_data = json.loads(result['json_output'])
                    big_chunks = json_data.get('big_chunks', [])
                    total_small_chunks = sum(len(chunk.get('small_chunks', [])) for chunk in big_chunks)
                    
                    # Content extraction metrics
                    st.markdown("#### Content Extraction")
                    colA, colB, colC = st.columns(3)
                    colA.metric("Big Chunks", len(big_chunks))
                    colB.metric("Total Small Chunks", total_small_chunks)
                    colC.metric("Content Length", f"{len(result['extracted_content']):,} chars")
                    
                    # Enhanced AI Analysis metrics
                    if 'ai_analysis_result' in st.session_state and st.session_state['ai_analysis_result'].get('success'):
                        st.markdown("#### AI Analysis Performance")
                        ai_result = st.session_state['ai_analysis_result']
                        
                        # Primary metrics
                        colD, colE, colF, colG = st.columns(4)
                        colD.metric("Processing Time", f"{ai_result.get('processing_time', 0):.2f}s")
                        colE.metric("Successful Analyses", ai_result.get('successful_count', 0))
                        colF.metric("Failed Analyses", ai_result.get('failed_count', 0), 
                                   delta=f"-{ai_result.get('failed_count', 0)}" if ai_result.get('failed_count', 0) > 0 else None)
                        colG.metric("Success Rate", f"{ai_result.get('success_rate', 0):.1f}%")
                        
                        # Secondary metrics
                        colH, colI, colJ, colK = st.columns(4)
                        colH.metric("Total Tokens", f"{ai_result.get('total_tokens', 0):,}")
                        colI.metric("Avg Time/Chunk", f"{ai_result.get('avg_processing_time', 0):.2f}s")
                        colJ.metric("Total Retries", ai_result.get('total_retries', 0))
                        
                        # Calculate cost estimate (approximate)
                        total_tokens = ai_result.get('total_tokens', 0)
                        if total_tokens > 0:
                            # Rough estimate: $0.01 per 1K tokens (varies by model)
                            estimated_cost = (total_tokens / 1000) * 0.01
                            colK.metric("Est. Cost", f"${estimated_cost:.3f}")
                        else:
                            colK.metric("Est. Cost", "N/A")
                        
                        # Performance insights
                        processing_time = ai_result.get('processing_time', 0)
                        total_chunks = ai_result.get('total_chunks', 0)
                        
                        if processing_time > 0 and total_chunks > 0:
                            throughput = total_chunks / processing_time
                            st.info(f"📊 **Throughput**: {throughput:.2f} chunks/second | **Parallel efficiency**: Achieved with max concurrency of {max_concurrent}")
                        
                        # Quality indicators
                        if ai_result.get('failed_count', 0) == 0:
                            st.success("🎯 **Quality**: Perfect success rate - all content analyzed successfully")
                        elif ai_result.get('success_rate', 0) >= 90:
                            st.info(f"✅ **Quality**: High success rate ({ai_result.get('success_rate', 0):.1f}%) - excellent analysis coverage")
                        elif ai_result.get('success_rate', 0) >= 70:
                            st.warning(f"⚠️ **Quality**: Moderate success rate ({ai_result.get('success_rate', 0):.1f}%) - some content may need manual review")
                        else:
                            st.error(f"❌ **Quality**: Low success rate ({ai_result.get('success_rate', 0):.1f}%) - significant issues detected")
                
                except (json.JSONDecodeError, TypeError):
                    st.warning("Could not parse JSON for statistics.")
                
                st.info(f"**Source URL**: {result['url']}")
        else:
            # Show original tabs without AI analysis
            tab1, tab2, tab3 = st.tabs(["🎯 JSON Output", "📄 Extracted Content", "📈 Summary"])
            
            with tab1:
                st.code(result['json_output'], language='json')
                st.download_button(
                    label="💾 Download JSON",
                    data=result['json_output'],
                    file_name=f"chunks_{int(time.time())}.json",
                    mime="application/json"
                )
            with tab2:
                st.text_area("Raw extracted content:", value=result['extracted_content'], height=400)
            with tab3:
                st.subheader("Processing Summary")
                try:
                    json_data = json.loads(result['json_output'])
                    big_chunks = json_data.get('big_chunks', [])
                    total_small_chunks = sum(len(chunk.get('small_chunks', [])) for chunk in big_chunks)
                    
                    colA, colB, colC = st.columns(3)
                    colA.metric("Big Chunks", len(big_chunks))
                    colB.metric("Total Small Chunks", total_small_chunks)
                    colC.metric("Content Length", f"{len(result['extracted_content']):,} chars")
                except (json.JSONDecodeError, TypeError):
                    st.warning("Could not parse JSON for statistics.")
                st.info(f"**Source URL**: {result['url']}")
        
        # Show API key reminder if not available
        if not api_key:
            st.info("💡 **Tip**: Add your OpenAI API key to enable enhanced AI compliance analysis!")
        
        # Additional help section
        with st.expander("💡 Tips for Better Results"):
            st.markdown("""
            **For optimal AI analysis performance:**
            
            🔧 **Configuration Tips:**
            - Start with Max Concurrent = 3 for most use cases
            - Increase Max Concurrent (up to 10) if you have higher rate limits
            - Use Max Retries = 3 for reliability vs speed balance
            - Increase Max Retries for unreliable network connections
            - Max Tokens: 30K is safe, reduce if content is too long
            
            ⚡ **Performance Optimization:**
            - Larger content = longer processing time (expected)
            - More chunks = better parallel efficiency
            - Monitor token usage to manage costs
            - Check success rates - 90%+ is excellent
            - Token estimation: ~4 characters = 1 token
            
            🛠️ **Troubleshooting:**
            - Rate limit errors: Reduce Max Concurrent
            - Timeout errors: Check internet connection, increase Max Retries
            - Content too long errors: Reduce Max Tokens setting
            - Assistant errors: Verify your OpenAI API key has access to Assistants
            - Token limit errors: Content will show estimated tokens in error message
            
            📊 **Understanding Results:**
            - Success Rate: Percentage of chunks successfully analyzed
            - Total Tokens: API usage (impacts costs)
            - Retry Count: Network reliability indicator
            - Processing Time: Total time including retries and delays
            """)

if __name__ == "__main__":
    main()
